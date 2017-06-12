# a pythonic screenplay formatter
# it's good


from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys, re, os


def open_read(path):
    return Document(path), path


def open_write():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Courier New'
    font.size = Pt(12)
    style_margins(doc,0,0,1.5,1)
    return doc


def split_bracketed(para,left,right):
    return [x for x in re.split(left+'|'+right,para) if x]


# accept transform here
def description(doc,text,keys,cast):
    text = transform(keys,text)
    for player in cast:
        for i in [x.start() for x in re.finditer(player,text.upper())]:
            text = text[:i]+player+text[len(player)+i:]
    style_paragraph(doc,text,0.0,0.0,None)


def is_subheader(header):
    return header == 'INT' or header == 'EXT' or header == 'SUB'


def heading(doc,header,text,keys):
    desc = transform(keys,text)
    heading = header+'. '+desc.strip().upper()
    style_paragraph(doc,heading,0.0,0.0,None)


def dialogue(doc,header,paragraph,tags,keys,cast):
    header = transform(tags,header)
    if header not in cast:
        cast.append(header)
    paragraph = transform(keys,paragraph.strip())
    style_paragraph(doc,header,2.2,2.0,0)
    if paragraph.startswith('('):
        delim = paragraph.strip('(').split(')',1)
        style_paragraph(doc,'('+delim[0]+')',1.6,2.0,0)
        paragraph = delim[1].strip()
    style_paragraph(doc,paragraph,1.0,1.5,None)


def transition(doc,text,keys):
    text = transform(keys,text)
    if not text.endswith(':'):
        text += ':'
    fmt = style_paragraph(doc,text,0.0,0.0,None)
    fmt.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_tags(trans,text):
    key,value = [x.strip().upper() for x in text.split('=')][:2]
    trans[key] = value


def add_keys(trans,text):
    key,value = [x.strip() for x in text.split('=')][:2]
    trans[key] = value


# keep an eye on this, make sure it continues to look good
def transform(trans,text):
    for key, value in trans.items():
        text = re.sub('((?<= )|^)'+key+'((?=[ .,?!])|$)',value,text)
    return text


def is_a_note(paragraph):
    text = paragraph.strip()
    return text[0] == '['


def style_margins(doc,top,bottom,left,right):
    for section in doc.sections:
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def style_paragraph(doc,text,left,right,carriage):
    fmt = doc.add_paragraph(text)
    fmt.paragraph_format.left_indent = Inches(left)
    fmt.paragraph_format.right_indent = Inches(right)
    fmt.style = doc.styles['Normal']
    if carriage is not None:
        fmt.paragraph_format.space_after = Pt(carriage)
    return fmt


def shape_entry(para_obj):
    try:
        paragraph = split_bracketed(para_obj.text,'<','>')
        header = paragraph[0].strip().upper()
    except IndexError:
        return None,None
    else:
        return paragraph,header


def save_doc(doc,path):
    name = path.split('.')
    saved_path = '.'.join(name[:-1])+'.formatted.'+name[-1]
    try:
        doc.save(saved_path)
    except PermissionError:
        print('Unable to save file. Try again after closing file with same name.')


def convert(path):
    read, path = open_read(path)
    write = open_write()
    tags,keys = {},{}
    cast = []
    for para_obj in read.paragraphs:
        paragraph,header = shape_entry(para_obj)
        if not paragraph:
            pass
        elif len(paragraph) == 1:
            if not is_a_note(paragraph[0]):
                description(write,paragraph[0],keys,cast)
        elif is_subheader(header):
            heading(write,header,paragraph[1],keys)
        elif header == 'TRAN':
            transition(write,paragraph[1].strip().upper(),keys)
        elif header == 'TAG':
            add_tags(tags,paragraph[1])
        elif header == 'KEY':
            add_keys(keys,paragraph[1])
        else:
            dialogue(write,header,paragraph[1],tags,keys,cast)
    save_doc(write,path)


def validity(path):

    def query_dir(path,dirc,file):
        if file in os.listdir(dirc):
            return path
        else:
            return print('That file does not exist.')

    if path == '':
        path = 'script.docx'
    elif not path.endswith('.docx'):
        path += '.docx'
    dirc = path.split('/',1)
    if len(dirc) > 1:
        return query_dir(path,dirc[0],dirc[1])
    else:
        return query_dir(path,None,path)


def help_prompt():
    help_msg = "\nThe current markup is '<' and '>'.\n" \
        "These can be changed by entering '--markup' and " \
        "the symbols separated by spaces.\n" \
        "If no file is entered, the program will look for script.docx" \
        "Enter 'exit' to quit.\n"
    print(help_msg)


def command_line_input(argv):
    path = validity(' '.join(sys.argv[1:]))
    if path:
        convert(path)    
    print('Process finished.')


def user_input():
    while True:
        prompt = "\nEnter filepath of .docx to format or " \
                "--help' for instructions.\n"
        path = input(prompt).strip()
        if path == 'exit':
            return print('Program ended')
        elif path == '--help':
            help_prompt()
        else:
            path = validity(path)
            if path:
                convert(path)


def driver():
    if len(sys.argv) > 1:
        command_line_input(sys.argv)
    else:
        user_input()


driver()


# A pythonic script that reads and formats scripts!

